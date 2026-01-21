#!/usr/bin/env python3
"""
Script to convert Markdown files to Word (.docx) format.
Handles tables, headers, lists, Mermaid diagrams, and basic formatting.
"""

import re
import subprocess
import tempfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path


def parse_markdown_table(lines):
    """Parse markdown table lines into a 2D array."""
    rows = []
    for line in lines:
        if line.strip().startswith('|') and not re.match(r'^\|[-:\s|]+\|$', line):
            cells = [cell.strip() for cell in line.strip().strip('|').split('|')]
            rows.append(cells)
    return rows


def add_table_to_doc(doc, rows):
    """Add a table to the document."""
    if not rows:
        return
    
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell
            # Bold header row
            if i == 0:
                for paragraph in table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def render_mermaid_to_image(mermaid_code, output_path):
    """Render Mermaid diagram to PNG image using mmdc CLI."""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False) as f:
        f.write(mermaid_code)
        mmd_file = f.name
    
    try:
        result = subprocess.run(
            ['mmdc', '-i', mmd_file, '-o', str(output_path), '-b', 'white', '-t', 'default'],
            capture_output=True,
            text=True,
            timeout=30
        )
        if result.returncode == 0 and Path(output_path).exists():
            return True
        else:
            print(f"Mermaid error: {result.stderr}")
            return False
    except Exception as e:
        print(f"Error rendering Mermaid: {e}")
        return False
    finally:
        Path(mmd_file).unlink(missing_ok=True)


def convert_md_to_docx(md_path, docx_path):
    """Convert a Markdown file to Word document."""
    doc = Document()
    
    # Read markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    mermaid_count = 0
    
    # Create temp directory for images
    temp_dir = Path(tempfile.gettempdir()) / 'mermaid_images'
    temp_dir.mkdir(exist_ok=True)
    
    while i < len(lines):
        line = lines[i]
        
        # Handle mermaid code blocks - render as image
        if line.strip().startswith('```mermaid'):
            i += 1
            mermaid_code = []
            while i < len(lines) and not lines[i].strip() == '```':
                mermaid_code.append(lines[i])
                i += 1
            i += 1
            
            # Render mermaid to image
            mermaid_count += 1
            img_path = temp_dir / f'mermaid_{mermaid_count}.png'
            mermaid_text = '\n'.join(mermaid_code)
            
            if render_mermaid_to_image(mermaid_text, img_path):
                # Add image to document
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(str(img_path), width=Inches(5.5))
                print(f"  ✓ Mermaid diagram {mermaid_count} rendered")
            else:
                doc.add_paragraph('[Diagramme Mermaid - erreur de rendu]', style='Intense Quote')
            continue
        
        # Skip other code blocks
        if line.strip().startswith('```'):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            heading = doc.add_heading(line[2:].strip(), level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=3)
        
        # Horizontal rule
        elif line.strip() == '---':
            doc.add_paragraph('─' * 50)
        
        # Tables
        elif line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table_to_doc(doc, parse_markdown_table(table_lines))
            continue
        
        # Blockquotes / Notes
        elif line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            # Remove markdown note syntax
            text = re.sub(r'\*\*Note\*\*\s*:', 'Note :', text)
            text = re.sub(r'\[!NOTE\]', '', text)
            text = re.sub(r'\[!IMPORTANT\]', '', text)
            if text:
                para = doc.add_paragraph(text, style='Intense Quote')
        
        # Standard Markdown Images
        elif re.match(r'^!\[(.*?)\]\((.*?)\)', line.strip()):
            match = re.match(r'^!\[(.*?)\]\((.*?)\)', line.strip())
            alt_text = match.group(1)
            img_src = match.group(2)
            
            # Resolve image path relative to markdown file
            img_path = Path(md_path).parent / img_src
            
            if img_path.exists():
                try:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    # Constrain width to page width (approx 6 inches)
                    run.add_picture(str(img_path), width=Inches(6))
                    # Add caption if alt text exists
                    if alt_text:
                        caption = doc.add_paragraph(alt_text, style='Caption')
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    print(f"  ✓ Image embedded: {img_src}")
                except Exception as e:
                    print(f"  ⚠️ Failed to embed image {img_src}: {e}")
                    doc.add_paragraph(f'[Image: {alt_text}]', style='Intense Quote')
            else:
                print(f"  ⚠️ Image not found: {img_path}")
                doc.add_paragraph(f'[Image manquante: {alt_text}]', style='Intense Quote')

        # Checkbox lists
        elif re.match(r'^-\s*\[([ x/])\]', line):
            match = re.match(r'^-\s*\[([ x/])\]\s*(.*)', line)
            if match:
                status = match.group(1)
                text = match.group(2)
                symbol = '☑' if status == 'x' else ('◐' if status == '/' else '☐')
                doc.add_paragraph(f'{symbol} {text}', style='List Bullet')
        
        # Regular lists
        elif line.strip().startswith('- '):
            text = line.strip()[2:]
            # Handle bold text
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')
        
        # Regular paragraphs
        elif line.strip():
            # Clean up markdown formatting
            text = line.strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Code
            if text:
                doc.add_paragraph(text)
        
        i += 1
    
    # Save document
    doc.save(docx_path)
    print(f'✅ Document créé : {docx_path}')


if __name__ == '__main__':
    import sys
    
    if len(sys.argv) < 2:
        # Default: convert cahier-des-charges.md
        md_file = Path(__file__).parent / 'requirements' / 'cahier-des-charges.md'
        docx_file = Path(__file__).parent / 'requirements' / 'cahier-des-charges.docx'
    else:
        md_file = Path(sys.argv[1])
        docx_file = md_file.with_suffix('.docx')
    
    convert_md_to_docx(md_file, docx_file)

